"""Read uploaded brand-voice documents and export generated content to PDF."""
from __future__ import annotations

import io
import re
from datetime import datetime

import markdown2
from xhtml2pdf import pisa


def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from an uploaded .txt/.md/.docx/.pdf file."""
    name = (filename or "").lower()
    if name.endswith((".txt", ".md", ".markdown")):
        return data.decode("utf-8", errors="replace")
    if name.endswith(".docx"):
        import docx  # python-docx
        doc = docx.Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if name.endswith(".pdf"):
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    # unknown extension: best-effort decode
    return data.decode("utf-8", errors="replace")


PDF_CSS = """
@page { size: A4; margin: 2.2cm 2cm; }
body { font-family: Helvetica, Arial, sans-serif; font-size: 10.5pt; line-height: 1.5; color: #1a1a1a; }
h1 { font-size: 20pt; color: #0f2b46; margin: 0 0 6pt 0; }
h2 { font-size: 14pt; color: #0f2b46; border-bottom: 1px solid #c8a24a; padding-bottom: 3pt; margin: 16pt 0 6pt; }
h3 { font-size: 12pt; color: #234; margin: 12pt 0 4pt; }
p, li { font-size: 10.5pt; }
table { border-collapse: collapse; width: 100%; margin: 8pt 0; }
th, td { border: 1px solid #bbb; padding: 3pt 4pt; text-align: left; font-size: 9pt; word-wrap: break-word; }
th { background: #0f2b46; color: #fff; }
blockquote { border-left: 3px solid #c8a24a; margin: 8pt 0; padding: 2pt 10pt; color: #444; font-style: italic; }
code { background: #f2f2f2; padding: 1pt 3pt; font-family: monospace; }
.cover { border-bottom: 2px solid #c8a24a; padding-bottom: 8pt; margin-bottom: 14pt; }
.cover .brand { color: #0f2b46; font-size: 22pt; font-weight: bold; }
.cover .sub { color: #666; font-size: 9pt; margin-top: 4pt; }
"""


def markdown_to_pdf(markdown_text: str, *, brand: str, topic: str,
                    article_type: str, language: str) -> bytes:
    """Render generated markdown content to PDF bytes."""
    body_html = markdown2.markdown(
        markdown_text,
        extras=["tables", "fenced-code-blocks", "cuddled-lists", "strike",
                "header-ids", "break-on-newline"],
    )
    stamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    cover = f"""
    <div class="cover">
      <div class="brand">{_esc(brand)}</div>
      <div class="sub">Project Kairos &middot; {_esc(article_type)} &middot;
      {_esc(language)} &middot; Generated {stamp}</div>
      <div class="sub"><b>Topic:</b> {_esc(topic)}</div>
    </div>
    """
    html = f"<html><head><style>{PDF_CSS}</style></head><body>{cover}{body_html}</body></html>"
    out = io.BytesIO()
    result = pisa.CreatePDF(src=html, dest=out, encoding="utf-8")
    if result.err:
        raise RuntimeError("PDF generation failed.")
    return out.getvalue()


def _fence(name: str) -> re.Pattern:
    return re.compile(rf"<<<\s*{name}_START\s*>>>(.*?)<<<\s*{name}_END\s*>>>",
                      re.DOTALL | re.IGNORECASE)


_PUBLISH_RE = _fence("PUBLISH_CONTENT")
_OPS_RE = _fence("OPS_PACK")
_SCORE_RE = _fence("SCORE_REPORT")

# internal grounding/citation tokens that must never survive into reader-facing content
_REF_TOKEN_RE = re.compile(
    r"\s*\[(?:graph:[^\]]+|web:[^\]]+|author-first-party[^\]]*|"
    r"not available[^\]]*|to verify[^\]]*|to source[^\]]*|source:[^\]]*)\]",
    re.IGNORECASE)


def sanitize_publish(text: str) -> str:
    """Make PART A truly copy-paste-ready even if the model slipped: strip any leftover
    internal reference tokens ([graph:...], [web:...], [graph:author-first-party-N],
    [to verify], ...) AND normalize em/en dashes to plain punctuation (KAIROS house style
    forbids em dashes). Reader content only - never applied to the score report."""
    if not text:
        return text
    out = _REF_TOKEN_RE.sub("", text)
    # house style: no em/en dashes. " - " keeps sentence flow; a bare dash becomes a hyphen.
    out = out.replace(" — ", " - ").replace(" – ", " - ")
    out = out.replace("—", "-").replace("–", "-")
    # tidy any spacing the removal left behind (double spaces, space-before-punctuation)
    out = re.sub(r"[ \t]{2,}", " ", out)
    out = re.sub(r" +([.,;:!?])", r"\1", out)
    return out


def split_sections(md: str) -> dict[str, str]:
    """Split model output into publish_content / ops_pack / score_report.

    Falls back gracefully if fences are missing. If no publish fence at all,
    the whole text is treated as publish_content. The publish content is sanitized
    of any internal citation tokens so it is genuinely paste-ready.
    """
    md = md or ""
    pub = _PUBLISH_RE.search(md)
    ops = _OPS_RE.search(md)
    score = _SCORE_RE.search(md)
    if not pub:
        # legacy / unfenced output: everything is the content
        return {"publish_content": sanitize_publish(md.strip()), "ops_pack": "", "score_report": ""}
    return {
        "publish_content": sanitize_publish(pub.group(1).strip()),
        "ops_pack": ops.group(1).strip() if ops else "",
        "score_report": score.group(1).strip() if score else "",
    }


def split_blocks(md: str) -> list[dict]:
    """Split publish content into logical blocks at heading boundaries.

    Each block = a heading line + its body up to the next heading. Content before
    the first heading becomes an 'Introduction' block. Returns [{index, heading, md}].
    """
    md = (md or "").strip()
    if not md:
        return []
    lines = md.split("\n")
    blocks: list[dict] = []
    cur: list[str] = []
    heading = "Introduction"

    def flush():
        text = "\n".join(cur).strip()
        if text:
            blocks.append({"index": len(blocks), "heading": heading, "md": text})

    for ln in lines:
        if re.match(r"^\s{0,3}#{1,4}\s+", ln):
            flush()
            cur = [ln]
            heading = re.sub(r"^\s{0,3}#{1,4}\s+", "", ln).strip()
        else:
            cur.append(ln)
    flush()
    for i, b in enumerate(blocks):
        b["index"] = i
    return blocks


def extract_fence(md: str, name: str) -> str:
    """Return the content between <<<NAME_START>>> and <<<NAME_END>>> (or '')."""
    m = _fence(name).search(md or "")
    return m.group(1).strip() if m else ""


def split_output(md: str) -> tuple[str, str]:
    """Back-compat: (publish_content, ops_pack+score_report combined)."""
    s = split_sections(md)
    pack = "\n\n".join(x for x in (s["ops_pack"], s["score_report"]) if x)
    return s["publish_content"], pack


def _esc(s: str) -> str:
    return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# --------------------------------------------------------------- DOCX export
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def _add_runs(paragraph, text: str) -> None:
    """Add text to a docx paragraph, honoring **bold**, *italic*, `code`."""
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Courier New"
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(markdown_text: str, *, brand: str, topic: str,
                     article_type: str, language: str) -> bytes:
    """Render generated markdown to a .docx (Word) file — good for CMS/editorial workflows."""
    import io as _io

    import docx
    from docx.shared import Pt, RGBColor

    doc = docx.Document()
    # cover line
    head = doc.add_paragraph()
    r = head.add_run(brand); r.bold = True; r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x0F, 0x2B, 0x46)
    sub = doc.add_paragraph()
    s = sub.add_run(f"{article_type} · {language}"); s.italic = True; s.font.size = Pt(9)

    lines = markdown_text.split("\n")
    i = 0
    num_re = re.compile(r"^\s*\d+\.\s+")
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        # table block
        if stripped.startswith("|") and i + 1 < len(lines) and set(lines[i + 1].strip()) <= set("|-: "):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if set("".join(cells)) <= set("-: "):
                    i += 1
                    continue
                rows.append(cells)
                i += 1
            if rows:
                ncol = max(len(r) for r in rows)
                table = doc.add_table(rows=0, cols=ncol)
                table.style = "Light Grid Accent 1"
                for ridx, cells in enumerate(rows):
                    tr = table.add_row().cells
                    for cidx in range(ncol):
                        txt = cells[cidx] if cidx < len(cells) else ""
                        p = tr[cidx].paragraphs[0]
                        _add_runs(p, txt)
                        if ridx == 0:
                            for run in p.runs:
                                run.bold = True
            continue
        # headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped in ("---", "***", "___"):
            pass
        elif stripped.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            _add_runs(p, stripped[2:])
        elif stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, stripped[2:])
        elif num_re.match(line):
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, num_re.sub("", line))
        else:
            p = doc.add_paragraph()
            _add_runs(p, stripped)
        i += 1

    out = _io.BytesIO()
    doc.save(out)
    return out.getvalue()


def safe_slug(text: str, maxlen: int = 60) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", (text or "content").lower()).strip("-")
    return (slug[:maxlen] or "content").strip("-")
